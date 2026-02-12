"""
Genera documento Word completo con:
1. Top 50 acciones semana actual por estrategia
2. Resultados CAGR y PnL por año
3. Lista completa de trades MIN_DRAWDOWN
"""
import psycopg2
from datetime import timedelta, datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

conn = psycopg2.connect('postgresql://fmp:fmp123@localhost:5433/fmp_data')
cur = conn.cursor()

COMMISSION_RATE = 0.003
POSITION_SIZE = 20000
MAX_POSITIONS = 50
MAX_CHANGES = 2

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Cargar semanas
cur.execute('''
    SELECT DISTINCT week_ending FROM market_cap_weekly
    WHERE week_ending >= '2023-01-01'
    ORDER BY week_ending
''')
weeks = [row[0] for row in cur.fetchall()]

FORMULAS = {
    'ORIGINAL': lambda peg, beat, eps, rev: (1.5 - peg) * 50 + min(eps, 150) * 0.4 + min(rev, 100) * 0.6 + min(beat, 15) * 2,
    'MIN_DRAWDOWN': lambda peg, beat, eps, rev: (1.0 - peg) * 60 + min(beat, 30) * 6 + min(eps, 80) * 0.2 + min(rev, 50) * 0.2,
    'MAX_RETURN': lambda peg, beat, eps, rev: (1.5 - peg) * 30 + min(eps, 200) * 0.8 + min(rev, 150) * 0.8 + min(beat, 10) * 1,
    'ORIG_TRANSCRIPT': lambda peg, beat, eps, rev: (1.5 - peg) * 40 + min(eps, 150) * 0.3 + min(rev, 100) * 0.5 + min(beat, 15) * 2,
}

def get_ranked_stocks(week_ending, formula_func):
    cur.execute('''
        SELECT m.symbol, b.beat_streak, p.peg_ratio, e.eps_growth_yoy, r.revenue_growth_yoy
        FROM market_cap_weekly m
        JOIN beat_streak_weekly b ON m.symbol = b.symbol AND m.week_ending = b.week_ending
        JOIN peg_weekly p ON m.symbol = p.symbol AND m.week_ending = p.week_ending
        JOIN eps_ttm_weekly e ON m.symbol = e.symbol AND m.week_ending = e.week_ending
        JOIN revenue_ttm_weekly r ON m.symbol = r.symbol AND m.week_ending = r.week_ending
        WHERE m.week_ending = %s
        AND m.market_cap >= 1000000000
        AND b.beat_streak >= 4
        AND p.peg_ratio > 0 AND p.peg_ratio <= 1.5
        AND e.eps_growth_yoy > 20
        AND r.revenue_growth_yoy > 12
        AND m.symbol NOT LIKE '%%.%%'
        AND m.symbol NOT LIKE '%%-%%'
        AND LENGTH(m.symbol) <= 5
        AND m.symbol !~ '[0-9]'
        AND RIGHT(m.symbol, 1) NOT IN ('F', 'Y')
    ''', (week_ending,))
    stocks = []
    for row in cur.fetchall():
        symbol, beat, peg, eps_g, rev_g = row
        peg = float(peg) if peg else 1.5
        beat = int(beat) if beat else 0
        eps_g = float(eps_g) if eps_g else 0
        rev_g = float(rev_g) if rev_g else 0
        stocks.append({'symbol': symbol, 'score': formula_func(peg, beat, eps_g, rev_g)})
    stocks.sort(key=lambda x: x['score'], reverse=True)
    return stocks

def get_monday_price(symbol, friday):
    monday = friday + timedelta(days=3)
    cur.execute('''
        SELECT date, close FROM fmp_price_history
        WHERE symbol = %s AND date >= %s AND date <= %s
        ORDER BY date LIMIT 1
    ''', (symbol, monday, monday + timedelta(days=5)))
    row = cur.fetchone()
    if row:
        return row[1], row[0]
    cur.execute('''
        SELECT date, close FROM fmp_price_history
        WHERE symbol = %s AND date <= %s ORDER BY date DESC LIMIT 1
    ''', (symbol, friday + timedelta(days=7)))
    row = cur.fetchone()
    return (row[1], row[0]) if row else (None, None)

def run_backtest(formula_func):
    positions = {}
    all_trades = []
    initial_capital = POSITION_SIZE * MAX_POSITIONS
    yearly_pnl = {}

    for i, week in enumerate(weeks[:-1]):
        ranked = get_ranked_stocks(week, formula_func)
        top50_symbols = set(s['symbol'] for s in ranked[:MAX_POSITIONS])
        ranked_dict = {s['symbol']: s['score'] for s in ranked}
        current_symbols = set(positions.keys())
        is_first_week = (i == 0)

        to_sell = current_symbols - top50_symbols
        to_buy = top50_symbols - current_symbols
        sells_list = sorted([(s, ranked_dict.get(s, -9999)) for s in to_sell], key=lambda x: x[1])
        buys_list = sorted([(s, ranked_dict.get(s, 0)) for s in to_buy], key=lambda x: x[1], reverse=True)

        max_sells = len(sells_list) if is_first_week else MAX_CHANGES
        sells_done = 0
        for symbol, _ in sells_list:
            if sells_done >= max_sells:
                break
            if symbol in positions:
                pos = positions[symbol]
                price, dt = get_monday_price(symbol, week)
                if price:
                    gross = price * pos['shares']
                    comm = gross * COMMISSION_RATE
                    pnl = gross - comm - pos['cost_basis']
                    pct = (pnl / pos['cost_basis']) * 100
                    duration = (dt - pos['entry_date']).days
                    duration_cagr = max(duration, 30)
                    cagr = ((1 + pct/100) ** (365 / duration_cagr) - 1) * 100
                    year = dt.year
                    yearly_pnl[year] = yearly_pnl.get(year, 0) + pnl
                    all_trades.append({
                        'symbol': symbol, 'entry_date': pos['entry_date'], 'exit_date': dt,
                        'duration': duration, 'pnl': pnl, 'pct': pct, 'cagr': cagr
                    })
                del positions[symbol]
                sells_done += 1

        max_buys = MAX_POSITIONS if is_first_week else MAX_CHANGES
        buys_done = 0
        for symbol, _ in buys_list:
            if buys_done >= max_buys or len(positions) >= MAX_POSITIONS:
                break
            price, dt = get_monday_price(symbol, week)
            if price and price > 0:
                shares = round(POSITION_SIZE / price)
                if shares > 0:
                    gross = price * shares
                    comm = gross * COMMISSION_RATE
                    positions[symbol] = {'shares': shares, 'entry_price': price, 'entry_date': dt, 'cost_basis': gross + comm}
                    buys_done += 1

    for symbol, pos in list(positions.items()):
        price, dt = get_monday_price(symbol, weeks[-2])
        if price:
            gross = price * pos['shares']
            comm = gross * COMMISSION_RATE
            pnl = gross - comm - pos['cost_basis']
            pct = (pnl / pos['cost_basis']) * 100
            duration = (dt - pos['entry_date']).days if dt else 0
            duration_cagr = max(duration, 30)
            cagr = ((1 + pct/100) ** (365 / duration_cagr) - 1) * 100
            year = dt.year if dt else weeks[-2].year
            yearly_pnl[year] = yearly_pnl.get(year, 0) + pnl
            all_trades.append({
                'symbol': symbol, 'entry_date': pos['entry_date'], 'exit_date': dt,
                'duration': duration, 'pnl': pnl, 'pct': pct, 'cagr': cagr
            })

    total_pnl = sum(t['pnl'] for t in all_trades)
    winners = [t for t in all_trades if t['pnl'] > 0]
    years_elapsed = (weeks[-2] - weeks[0]).days / 365.25
    portfolio_cagr = ((initial_capital + total_pnl) / initial_capital) ** (1/years_elapsed) - 1 if years_elapsed > 0 else 0

    return {
        'trades': all_trades,
        'total_pnl': total_pnl,
        'winners': len(winners),
        'total': len(all_trades),
        'win_rate': len(winners)/len(all_trades)*100 if all_trades else 0,
        'yearly_pnl': yearly_pnl,
        'portfolio_cagr': portfolio_cagr * 100
    }

# ============================================================================
# CREAR DOCUMENTO
# ============================================================================
print('Creando documento Word...', flush=True)
doc = Document()

# Titulo
title = doc.add_heading('Backtest PatrimonioSmart - Informe Completo', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'Generado: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
doc.add_paragraph(f'Periodo: {weeks[0]} a {weeks[-1]} ({len(weeks)} semanas)')
doc.add_paragraph('Capital inicial: $1,000,000 (50 posiciones x $20,000)')
doc.add_paragraph()

# ============================================================================
# SECCION 1: TOP 50 SEMANA ACTUAL
# ============================================================================
print('Generando Top 50 semana actual...', flush=True)
doc.add_heading('1. Top 50 Acciones - Semana 2026-02-06', level=1)

week_actual = '2026-02-06'
rankings = {}
for name, func in FORMULAS.items():
    stocks = get_ranked_stocks(week_actual, func)
    rankings[name] = [s['symbol'] for s in stocks[:50]]

table = doc.add_table(rows=51, cols=5)
table.style = 'Table Grid'

header = table.rows[0].cells
headers = ['#', 'ORIGINAL', 'MIN_DRAWDOWN', 'MAX_RETURN', 'ORIG_TRANSCRIPT']
for i, h in enumerate(headers):
    header[i].text = h
    header[i].paragraphs[0].runs[0].bold = True
    set_cell_shading(header[i], '4472C4')
    header[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

for i in range(50):
    row = table.rows[i+1].cells
    row[0].text = str(i+1)
    row[1].text = rankings['ORIGINAL'][i] if i < len(rankings['ORIGINAL']) else ''
    row[2].text = rankings['MIN_DRAWDOWN'][i] if i < len(rankings['MIN_DRAWDOWN']) else ''
    row[3].text = rankings['MAX_RETURN'][i] if i < len(rankings['MAX_RETURN']) else ''
    row[4].text = rankings['ORIG_TRANSCRIPT'][i] if i < len(rankings['ORIG_TRANSCRIPT']) else ''

doc.add_paragraph()

# ============================================================================
# SECCION 2: RESULTADOS POR ESTRATEGIA
# ============================================================================
doc.add_heading('2. Resultados del Backtest por Estrategia', level=1)

results = {}
for name, func in FORMULAS.items():
    print(f'  Ejecutando backtest {name}...', flush=True)
    results[name] = run_backtest(func)

# Tabla resumen
doc.add_heading('Resumen de Metricas', level=2)
summary_table = doc.add_table(rows=5, cols=8)
summary_table.style = 'Table Grid'

header = summary_table.rows[0].cells
for i, h in enumerate(['Estrategia', 'CAGR', 'Total PnL', 'Trades', 'Win Rate', 'PnL 2024', 'PnL 2025', 'PnL 2026']):
    header[i].text = h
    header[i].paragraphs[0].runs[0].bold = True
    set_cell_shading(header[i], '4472C4')
    header[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

for i, (name, r) in enumerate(sorted(results.items(), key=lambda x: x[1]['portfolio_cagr'], reverse=True)):
    row = summary_table.rows[i+1].cells
    row[0].text = name
    row[1].text = f"{r['portfolio_cagr']:.1f}%"
    row[2].text = f"${r['total_pnl']:,.0f}"
    row[3].text = str(r['total'])
    row[4].text = f"{r['win_rate']:.1f}%"
    row[5].text = f"${r['yearly_pnl'].get(2024, 0):,.0f}"
    row[6].text = f"${r['yearly_pnl'].get(2025, 0):,.0f}"
    row[7].text = f"${r['yearly_pnl'].get(2026, 0):,.0f}"

doc.add_paragraph()

# ============================================================================
# SECCION 3: LISTA COMPLETA DE TRADES (MIN_DRAWDOWN)
# ============================================================================
print('Generando lista de trades MIN_DRAWDOWN...', flush=True)
doc.add_heading('3. Lista Completa de Trades - MIN_DRAWDOWN', level=1)

trades = sorted(results['MIN_DRAWDOWN']['trades'], key=lambda x: x['entry_date'])
r = results['MIN_DRAWDOWN']
doc.add_paragraph(f"Total: {r['total']} trades | PnL: ${r['total_pnl']:,.0f} | Ganadores: {r['winners']} ({r['win_rate']:.1f}%)")

# Crear tabla de trades
trade_table = doc.add_table(rows=len(trades)+1, cols=7)
trade_table.style = 'Table Grid'

header = trade_table.rows[0].cells
for i, h in enumerate(['#', 'Ticker', 'Entrada', 'Salida', 'PnL $', 'Ret %', 'CAGR %']):
    header[i].text = h
    header[i].paragraphs[0].runs[0].bold = True
    set_cell_shading(header[i], '2E7D32')
    header[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

for j, t in enumerate(trades):
    row = trade_table.rows[j+1].cells
    row[0].text = str(j + 1)
    row[1].text = t['symbol']
    row[2].text = str(t['entry_date'])
    row[3].text = str(t['exit_date'])
    row[4].text = f"${t['pnl']:+,.0f}"
    row[5].text = f"{t['pct']:+.1f}%"
    row[6].text = f"{t['cagr']:+.1f}%"

    # Colorear perdidas en rojo
    if t['pnl'] < 0:
        for cell in [row[4], row[5], row[6]]:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()

# ============================================================================
# GUARDAR
# ============================================================================
output_path = 'C:/Users/usuario/financial-data-project/backtest_report_completo.docx'
doc.save(output_path)
print(f'\nDocumento guardado en: {output_path}', flush=True)

cur.close()
conn.close()
