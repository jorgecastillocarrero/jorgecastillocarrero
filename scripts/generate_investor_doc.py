"""
Generador de documento para presentación a inversores
PatrimonioSmart - Sistema de Gestión de Carteras Inteligente
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os

def create_investor_document():
    doc = Document()

    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ========== PORTADA ==========
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title_run = title.add_run("PATRIMONIOSMART")
    title_run.bold = True
    title_run.font.size = Pt(36)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("Sistema Inteligente de Gestión de Carteras")
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = RGBColor(80, 80, 80)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tag_run = tagline.add_run("Inversión Cuantitativa Adaptativa con Inteligencia Artificial")
    tag_run.italic = True
    tag_run.font.size = Pt(14)
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Documento Confidencial - {datetime.now().strftime('%B %Y')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ========== ÍNDICE ==========
    doc.add_heading('Contenido', level=1)

    toc_items = [
        "1. Resumen Ejecutivo",
        "2. El Problema: Limitaciones de la Inversión Tradicional",
        "3. Nuestra Solución: Inversión Cuantitativa Adaptativa",
        "4. Arquitectura Tecnológica",
        "5. Motor de Análisis Multi-Factor",
        "6. Regímenes de Mercado y Adaptabilidad",
        "7. Casos de Uso Prácticos",
        "   7.1 Top 50 Empresas Óptimas (Cartera Core)",
        "   7.2 Top 5 S&P 500 Swing Trading (5 días)",
        "8. Infraestructura de Datos",
        "9. Roadmap y Escalabilidad",
        "10. Motor Macroeconómico",
        "11. Ventajas Competitivas"
    ]

    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

    # ========== 1. RESUMEN EJECUTIVO ==========
    doc.add_heading('1. Resumen Ejecutivo', level=1)

    doc.add_paragraph(
        "PatrimonioSmart es una plataforma de gestión patrimonial de nueva generación que "
        "combina análisis cuantitativo avanzado, inteligencia artificial y procesamiento de "
        "lenguaje natural para construir y gestionar carteras de inversión óptimas."
    )

    doc.add_paragraph(
        "A diferencia de los sistemas tradicionales que se basan únicamente en métricas "
        "financieras estáticas, nuestra plataforma integra múltiples fuentes de información "
        "para adaptarse dinámicamente a las condiciones del mercado:"
    )

    bullets = [
        "Análisis fundamental y valoración cuantitativa (P/E, EV/EBITDA, ROE, márgenes)",
        "Indicadores técnicos y de momentum (61 métricas calculadas diariamente)",
        "Análisis de sentimiento de noticias y earnings calls mediante NLP",
        "Detección automática de regímenes de mercado (bull, bear, alta volatilidad)",
        "Ratios diana adaptativos que se ajustan según el contexto macroeconómico"
    ]

    for bullet in bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')

    # Caja de métricas clave
    doc.add_paragraph()
    metrics_table = doc.add_table(rows=2, cols=4)
    metrics_table.style = 'Table Grid'

    headers = ['Símbolos', 'Registros', 'Datos Históricos', 'Actualizaciones']
    values = ['+86,000', '+100 Millones', '20+ años', 'Diarias']

    for i, header in enumerate(headers):
        cell = metrics_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, value in enumerate(values):
        cell = metrics_table.rows[1].cells[i]
        cell.text = value
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ========== 2. EL PROBLEMA ==========
    doc.add_heading('2. El Problema: Limitaciones de la Inversión Tradicional', level=1)

    doc.add_heading('Enfoque Estático', level=2)
    doc.add_paragraph(
        "Los sistemas de inversión tradicionales aplican las mismas reglas independientemente "
        "del contexto de mercado. Una estrategia value que funciona en mercados alcistas puede "
        "ser desastrosa en periodos de alta volatilidad o crisis."
    )

    doc.add_heading('Datos Fragmentados', level=2)
    doc.add_paragraph(
        "Los inversores suelen analizar precio, fundamentales y noticias de forma aislada. "
        "La información no está integrada, lo que dificulta tomar decisiones holísticas."
    )

    doc.add_heading('Sesgo Emocional', level=2)
    doc.add_paragraph(
        "Las decisiones humanas están sujetas a sesgos cognitivos (aversión a pérdidas, "
        "efecto manada, exceso de confianza) que deterioran los retornos a largo plazo."
    )

    doc.add_heading('Información Desestructurada', level=2)
    doc.add_paragraph(
        "El 80% de la información relevante (noticias, transcripts de earnings, reportes) "
        "está en formato texto. Sin NLP, esta información no se puede procesar a escala."
    )

    # ========== 3. NUESTRA SOLUCIÓN ==========
    doc.add_heading('3. Nuestra Solución: Inversión Cuantitativa Adaptativa', level=1)

    doc.add_paragraph(
        "PatrimonioSmart resuelve estos problemas mediante un enfoque de inversión "
        "cuantitativa adaptativa que:"
    )

    solutions = [
        ("Integra Múltiples Fuentes",
         "Combina datos de precio, fundamentales, sentimiento y macro en un único modelo decisional."),
        ("Se Adapta al Régimen de Mercado",
         "Detecta automáticamente si estamos en bull market, bear market o alta volatilidad, "
         "ajustando los parámetros de la estrategia."),
        ("Usa Ratios Diana Dinámicos",
         "Los objetivos de P/E, volatilidad y exposición sectorial se adaptan según las "
         "condiciones macroeconómicas."),
        ("Procesa Información No Estructurada",
         "Mediante modelos de NLP (FinBERT, RoBERTa), analiza noticias y transcripts de "
         "earnings para extraer señales de sentimiento."),
        ("Elimina el Sesgo Emocional",
         "Las decisiones se basan en datos y reglas predefinidas, no en emociones.")
    ]

    for title, desc in solutions:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ========== 4. ARQUITECTURA TECNOLÓGICA ==========
    doc.add_heading('4. Arquitectura Tecnológica', level=1)

    doc.add_paragraph(
        "El sistema está construido sobre una arquitectura moderna y escalable, diseñada "
        "para procesar grandes volúmenes de datos financieros en tiempo real."
    )

    doc.add_heading('Stack Tecnológico', level=2)

    tech_table = doc.add_table(rows=8, cols=2)
    tech_table.style = 'Table Grid'

    tech_data = [
        ('Componente', 'Tecnología'),
        ('Lenguaje Principal', 'Python 3.12'),
        ('Base de Datos', 'PostgreSQL (escalable a cluster)'),
        ('Procesamiento de Datos', 'Pandas, NumPy, SQLAlchemy'),
        ('Machine Learning', 'Scikit-learn, XGBoost, LightGBM'),
        ('Deep Learning / NLP', 'PyTorch, Transformers, FinBERT'),
        ('Visualización', 'Streamlit (Dashboard interactivo)'),
        ('Infraestructura', 'Docker, Cloud-ready')
    ]

    for i, (comp, tech) in enumerate(tech_data):
        tech_table.rows[i].cells[0].text = comp
        tech_table.rows[i].cells[1].text = tech
        if i == 0:
            tech_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            tech_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    doc.add_heading('Flujo de Datos', level=2)

    flow_steps = [
        "1. Ingesta: Descarga diaria de precios, fundamentales, noticias y earnings",
        "2. Procesamiento: Cálculo de 130+ features técnicas, momentum y fundamentales",
        "3. Análisis NLP: Extracción de sentimiento de textos financieros",
        "4. Detección de Régimen: Clasificación automática del estado del mercado",
        "5. Optimización: Selección de activos según ratios diana adaptativos",
        "6. Ejecución: Generación de señales de compra/venta",
        "7. Monitorización: Dashboard en tiempo real con alertas"
    ]

    for step in flow_steps:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_page_break()

    # ========== 5. MOTOR DE ANÁLISIS MULTI-FACTOR ==========
    doc.add_heading('5. Motor de Análisis Multi-Factor', level=1)

    doc.add_paragraph(
        "El corazón del sistema es un motor de análisis que combina múltiples factores "
        "para evaluar cada activo. Cada día se calculan más de 130 métricas por símbolo."
    )

    doc.add_heading('5.1 Features Técnicas (61 métricas)', level=2)

    technical_categories = [
        ("Tendencia", "Medias móviles (SMA 20/50/200, EMA 12/26), ADX, dirección de tendencia"),
        ("Momentum", "RSI, RSI(2), MACD, Estocástico, Williams %R, CCI, MFI"),
        ("Volatilidad", "ATR, Bandas de Bollinger, Canales de Keltner, volatilidad histórica"),
        ("Volumen", "Ratio de volumen, OBV, Chaikin Money Flow, Accumulation/Distribution"),
        ("Señales", "Parabolic SAR, Ichimoku, Squeeze (Bollinger + Keltner), Donchian Breakout"),
        ("Price Action", "Gaps, nuevos máximos/mínimos 52 semanas, drawdown")
    ]

    for cat, desc in technical_categories:
        p = doc.add_paragraph()
        run = p.add_run(f"{cat}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading('5.2 Features de Momentum (25 métricas)', level=2)

    momentum_items = [
        "Retornos históricos: 1d, 5d, 20d, 60d, 252d",
        "Retornos futuros (target para ML): 1d, 5d, 20d forward",
        "Métricas de riesgo: Sharpe Ratio, Sortino Ratio, Max Drawdown",
        "Betas rolling vs S&P 500: 20d, 60d, 120d, 252d",
        "Momentum Score compuesto"
    ]

    for item in momentum_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('5.3 Features Fundamentales (34 métricas)', level=2)

    fundamental_items = [
        "Valoración: P/E, P/E Forward, P/B, P/S, EV/EBITDA",
        "Crecimiento: Revenue Growth 3Y/5Y, EPS Growth 3Y/5Y",
        "Rentabilidad: ROE, ROA, Gross Margin, Operating Margin, Profit Margin",
        "Salud financiera: Debt/Equity, Current Ratio, Total Debt",
        "Clasificación: Sector, Industria, Market Cap Category, Índices (S&P 500, Nasdaq 100)"
    ]

    for item in fundamental_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('5.4 Features de Earnings (10 métricas)', level=2)

    doc.add_paragraph(
        "Seguimiento de earnings surprises, días hasta próximo earnings, "
        "historial de beats consecutivos, y revenue surprises promedio."
    )

    doc.add_heading('5.5 Features de Sentimiento (NLP)', level=2)

    doc.add_paragraph(
        "Análisis de sentimiento de noticias financieras y transcripts de earnings calls "
        "utilizando modelos de lenguaje especializados en finanzas (FinBERT). "
        "Se calcula un score de sentimiento diario por activo, con ponderación temporal "
        "que da más peso a noticias recientes."
    )

    doc.add_page_break()

    # ========== 6. REGÍMENES DE MERCADO ==========
    doc.add_heading('6. Regímenes de Mercado y Adaptabilidad', level=1)

    doc.add_paragraph(
        "Una de las innovaciones clave del sistema es la detección automática de regímenes "
        "de mercado y la adaptación dinámica de los parámetros de inversión."
    )

    doc.add_heading('6.1 Detección de Regímenes', level=2)

    doc.add_paragraph(
        "El sistema analiza múltiples indicadores macroeconómicos y de mercado para "
        "clasificar el estado actual en uno de los siguientes regímenes:"
    )

    regimes_table = doc.add_table(rows=5, cols=3)
    regimes_table.style = 'Table Grid'

    regimes_data = [
        ('Régimen', 'Características', 'Estrategia'),
        ('Bull Market', 'Tendencia alcista, baja volatilidad, sentimiento positivo',
         'Mayor exposición a growth y momentum'),
        ('Bear Market', 'Tendencia bajista, VIX elevado, sentimiento negativo',
         'Reducir exposición, priorizar quality y defensivos'),
        ('Alta Volatilidad', 'VIX > 25, movimientos bruscos, incertidumbre',
         'Reducir posiciones, aumentar cash, stop-loss ajustados'),
        ('Lateral / Rango', 'Sin tendencia clara, baja volatilidad',
         'Estrategias mean reversion, opciones')
    ]

    for i, row_data in enumerate(regimes_data):
        for j, cell_text in enumerate(row_data):
            regimes_table.rows[i].cells[j].text = cell_text
            if i == 0:
                regimes_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    doc.add_heading('6.2 Ratios Diana Adaptativos', level=2)

    doc.add_paragraph(
        "Los parámetros objetivo de la cartera se ajustan dinámicamente según el régimen:"
    )

    ratios_items = [
        "P/E Objetivo: Más restrictivo en bear markets (< 15), más flexible en bull markets (< 25)",
        "Volatilidad Máxima: Reducida en regímenes de alta volatilidad",
        "Beta Objetivo: < 1 en bear markets, > 1 en bull markets para capturar momentum",
        "Exposición Sectorial: Rotar hacia defensivos (utilities, healthcare) en bear markets",
        "Nivel de Cash: Aumentar en periodos de incertidumbre",
        "Concentración: Más diversificada en alta volatilidad, más concentrada con convicción en bull"
    ]

    for item in ratios_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('6.3 Señales Macro Incorporadas', level=2)

    macro_signals = [
        "VIX (Índice de Volatilidad): Mide el miedo del mercado",
        "Curva de Tipos: Inversión de curva como predictor de recesión",
        "Breadth del Mercado: % de acciones sobre SMA 200",
        "Put/Call Ratio: Sentimiento de opciones",
        "Flujos de Fondos: Entradas/salidas de ETFs principales"
    ]

    for signal in macro_signals:
        doc.add_paragraph(signal, style='List Bullet')

    doc.add_page_break()

    # ========== 7. CASO DE USO: TOP 50 EMPRESAS ==========
    doc.add_heading('7. Caso de Uso: Top 50 Empresas Óptimas', level=1)

    doc.add_paragraph(
        "El siguiente ejemplo ilustra cómo el sistema genera una lista de las 50 mejores "
        "empresas para invertir, aplicando filtros predefinidos y adaptándose al contexto "
        "actual del mercado."
    )

    doc.add_heading('7.1 El Problema que Resuelve', level=2)

    doc.add_paragraph(
        "Un inversor quiere construir una cartera concentrada de calidad con las siguientes "
        "restricciones:"
    )

    problem_items = [
        "Seleccionar las 50 mejores empresas del universo de +86,000 símbolos",
        "Aplicar filtros cuantitativos estrictos (valoración, calidad, momentum)",
        "Limitar a máximo 2 operaciones de entrada/salida por periodo para reducir costes y rotación",
        "Adaptar los criterios de selección al régimen de mercado actual",
        "Maximizar los ratios diana (Sharpe, Sortino) ajustados por el contexto macro"
    ]

    for item in problem_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7.2 Pipeline de Selección', level=2)

    doc.add_paragraph("El sistema ejecuta un pipeline de 5 etapas:")

    # Etapa 1
    p = doc.add_paragraph()
    run = p.add_run("Etapa 1 - Filtros de Liquidez y Calidad Base: ")
    run.bold = True
    p.add_run("De 86,000 símbolos → 5,000 candidatos")

    filter1_items = [
        "Market Cap > $2B (elimina micro y small caps ilíquidos)",
        "Volumen medio diario > $10M",
        "Sin penny stocks (precio > $5)",
        "Datos fundamentales disponibles (excluye SPACs, shells)"
    ]
    for item in filter1_items:
        doc.add_paragraph(item, style='List Bullet')

    # Etapa 2
    p = doc.add_paragraph()
    run = p.add_run("Etapa 2 - Filtros Fundamentales: ")
    run.bold = True
    p.add_run("De 5,000 → 500 candidatos")

    filter2_items = [
        "ROE > 12% (rentabilidad del capital)",
        "Profit Margin > 8% (empresas rentables)",
        "Debt/Equity < 1.5 (salud financiera)",
        "Current Ratio > 1.2 (liquidez a corto plazo)",
        "Revenue Growth 3Y > 5% (crecimiento sostenido)"
    ]
    for item in filter2_items:
        doc.add_paragraph(item, style='List Bullet')

    # Etapa 3
    p = doc.add_paragraph()
    run = p.add_run("Etapa 3 - Filtros Técnicos y Momentum: ")
    run.bold = True
    p.add_run("De 500 → 150 candidatos")

    filter3_items = [
        "Precio > SMA 200 (tendencia alcista de largo plazo)",
        "RSI entre 30-70 (ni sobrecomprado ni sobrevendido extremo)",
        "Momentum Score > 0 (momentum positivo)",
        "No en zona de 'crash' (caída > 20% en 20 días)"
    ]
    for item in filter3_items:
        doc.add_paragraph(item, style='List Bullet')

    # Etapa 4
    p = doc.add_paragraph()
    run = p.add_run("Etapa 4 - Incorporación de Información Actual del Mercado: ")
    run.bold = True
    p.add_run("Ajuste dinámico de ratios diana")

    doc.add_paragraph(
        "En esta etapa crítica, el sistema analiza el contexto actual del mercado y ajusta "
        "los parámetros de selección:"
    )

    # Tabla de ajustes por régimen
    regime_adjust_table = doc.add_table(rows=5, cols=4)
    regime_adjust_table.style = 'Table Grid'

    regime_adjust_data = [
        ('Parámetro', 'Bull Market', 'Bear Market', 'Alta Volatilidad'),
        ('P/E Máximo', '< 30', '< 18', '< 20'),
        ('Beta Objetivo', '1.0 - 1.3', '0.6 - 0.9', '0.5 - 0.8'),
        ('Sharpe Mínimo', '> 0.8', '> 1.2', '> 1.5'),
        ('% Defensivos', '20%', '50%', '40%')
    ]

    for i, row_data in enumerate(regime_adjust_data):
        for j, cell_text in enumerate(row_data):
            regime_adjust_table.rows[i].cells[j].text = cell_text
            if i == 0 or j == 0:
                regime_adjust_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Etapa 5
    p = doc.add_paragraph()
    run = p.add_run("Etapa 5 - Ranking y Selección Final: ")
    run.bold = True
    p.add_run("De 150 → 50 empresas")

    doc.add_paragraph(
        "Se calcula un score compuesto ponderado por el régimen actual:"
    )

    score_items = [
        "Calidad Fundamental (25%): ROE, márgenes, crecimiento",
        "Valoración Relativa (20%): P/E vs sector, P/B vs histórico",
        "Momentum Técnico (20%): Fuerza relativa, tendencia, volumen",
        "Sentimiento NLP (15%): Score de noticias y earnings calls",
        "Riesgo Ajustado (20%): Sharpe, Sortino, Beta, Max Drawdown"
    ]
    for item in score_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7.3 Control de Rotación: Máximo 2 Operaciones', level=2)

    doc.add_paragraph(
        "Para minimizar costes de transacción y evitar sobre-trading, el sistema implementa "
        "un control de rotación inteligente:"
    )

    rotation_items = [
        "Umbral de entrada: Una acción solo entra al Top 50 si su score supera el percentil 95",
        "Umbral de salida: Una acción solo sale si cae por debajo del percentil 80",
        "Histéresis: Esta banda de histéresis evita rotación excesiva por pequeñas fluctuaciones",
        "Límite de operaciones: Máximo 2 cambios (1 entrada + 1 salida, o 2 entradas, o 2 salidas) por periodo de rebalanceo",
        "Cola de prioridad: Si hay más de 2 señales, se ejecutan las de mayor impacto en el score de cartera"
    ]
    for item in rotation_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7.4 Ratios Diana Adaptativos', level=2)

    doc.add_paragraph(
        "Los objetivos de la cartera se ajustan automáticamente según las condiciones del mercado:"
    )

    # Tabla de ratios diana
    diana_table = doc.add_table(rows=6, cols=3)
    diana_table.style = 'Table Grid'

    diana_data = [
        ('Ratio Diana', 'Valor Base', 'Ajuste Dinámico'),
        ('Sharpe Ratio Objetivo', '> 1.0', '+0.3 en alta volatilidad'),
        ('Max Drawdown Límite', '< 15%', 'Reducir a 10% si VIX > 25'),
        ('Concentración Top 10', '40%', 'Reducir a 30% en bear market'),
        ('Exposición Sectorial Max', '25%', 'Aumentar defensivos en bear'),
        ('Nivel de Cash', '5%', 'Hasta 20% en alta volatilidad')
    ]

    for i, row_data in enumerate(diana_data):
        for j, cell_text in enumerate(row_data):
            diana_table.rows[i].cells[j].text = cell_text
            if i == 0:
                diana_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    doc.add_heading('7.5 Ejemplo de Output', level=2)

    doc.add_paragraph(
        "El sistema genera un informe diario con:"
    )

    output_items = [
        "Lista ordenada de las 50 empresas con score, sector y métricas clave",
        "Señales de entrada/salida (máximo 2 por periodo)",
        "Régimen de mercado detectado y justificación",
        "Ratios diana actuales y desviación de la cartera",
        "Alertas: earnings próximos, noticias relevantes, cambios de rating"
    ]
    for item in output_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        "Este enfoque sistemático permite gestionar una cartera de alta calidad con mínima "
        "rotación, adaptándose automáticamente a las condiciones cambiantes del mercado."
    )

    doc.add_page_break()

    # ========== 7B. CASO DE USO: SWING TRADING 5 DÍAS ==========
    doc.add_heading('7B. Caso de Uso: Top 5 Acciones S&P 500 para Swing Trading', level=1)

    doc.add_paragraph(
        "Este caso de uso ilustra cómo el sistema identifica las 5 mejores oportunidades "
        "del S&P 500 para una operación táctica de corto plazo (5 días), combinando "
        "análisis técnico avanzado con el contexto actual del mercado."
    )

    doc.add_heading('7B.1 Parámetros de la Operación', level=2)

    swing_params = [
        "Universo: S&P 500 (500 empresas de alta liquidez)",
        "Horizonte: 5 días de trading (entrada al cierre, salida al cierre del día 5)",
        "Número de posiciones: 5 acciones",
        "Criterios: Técnicos + contexto de mercado",
        "Objetivo: Maximizar probabilidad de retorno positivo en 5 días"
    ]
    for item in swing_params:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7B.2 Análisis del Contexto de Mercado', level=2)

    doc.add_paragraph(
        "Antes de seleccionar acciones individuales, el sistema evalúa el estado del mercado:"
    )

    # Tabla de indicadores macro
    macro_table = doc.add_table(rows=6, cols=3)
    macro_table.style = 'Table Grid'

    macro_data = [
        ('Indicador', 'Lectura', 'Interpretación'),
        ('VIX', '< 20', 'Baja volatilidad, entorno favorable'),
        ('S&P 500 vs SMA 50', 'Por encima', 'Tendencia de corto plazo alcista'),
        ('% Acciones > SMA 200', '> 60%', 'Breadth saludable'),
        ('Put/Call Ratio', '< 0.8', 'Sentimiento no excesivamente bajista'),
        ('RSI del SPY', '40-60', 'No sobrecomprado ni sobrevendido')
    ]

    for i, row_data in enumerate(macro_data):
        for j, cell_text in enumerate(row_data):
            macro_table.rows[i].cells[j].text = cell_text
            if i == 0:
                macro_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    doc.add_heading('7B.3 Filtros de Selección Técnica', level=2)

    doc.add_paragraph("El sistema aplica una cascada de filtros técnicos:")

    # Filtro 1: Tendencia
    p = doc.add_paragraph()
    run = p.add_run("1. Filtro de Tendencia Favorable: ")
    run.bold = True

    trend_filters = [
        "Precio > SMA 50 (tendencia de medio plazo alcista)",
        "Precio > SMA 200 (tendencia de largo plazo alcista)",
        "SMA 50 > SMA 200 (golden cross activo)",
        "ADX > 20 (tendencia definida, no lateral)"
    ]
    for item in trend_filters:
        doc.add_paragraph(item, style='List Bullet')

    # Filtro 2: Momentum
    p = doc.add_paragraph()
    run = p.add_run("2. Filtro de Momentum Óptimo: ")
    run.bold = True

    momentum_filters = [
        "RSI(14) entre 40-65 (ni sobrevendido extremo ni sobrecomprado)",
        "RSI(2) < 20 (pullback de corto plazo para entry)",
        "MACD > Signal Line (momentum positivo)",
        "Momentum Score en percentil > 70"
    ]
    for item in momentum_filters:
        doc.add_paragraph(item, style='List Bullet')

    # Filtro 3: Volumen
    p = doc.add_paragraph()
    run = p.add_run("3. Filtro de Volumen y Liquidez: ")
    run.bold = True

    volume_filters = [
        "Volumen ratio > 0.8 (volumen no anormalmente bajo)",
        "Chaikin Money Flow > 0 (dinero entrando)",
        "OBV slope positivo (acumulación)"
    ]
    for item in volume_filters:
        doc.add_paragraph(item, style='List Bullet')

    # Filtro 4: Volatilidad
    p = doc.add_paragraph()
    run = p.add_run("4. Filtro de Volatilidad Controlada: ")
    run.bold = True

    vol_filters = [
        "ATR% < 3% (volatilidad diaria manejable)",
        "No en squeeze extremo (Bollinger Width normal)",
        "Keltner position entre -0.5 y 0.5 (no en extremos)"
    ]
    for item in vol_filters:
        doc.add_paragraph(item, style='List Bullet')

    # Filtro 5: Eventos
    p = doc.add_paragraph()
    run = p.add_run("5. Filtro de Eventos: ")
    run.bold = True

    event_filters = [
        "Sin earnings en los próximos 5 días (evitar saltos impredecibles)",
        "Sin ex-dividend date en el periodo",
        "Sin noticias de alto impacto pendientes (FDA, antitrust, etc.)"
    ]
    for item in event_filters:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7B.4 Ranking y Selección Final', level=2)

    doc.add_paragraph(
        "De las acciones que pasan todos los filtros, se calcula un score de oportunidad:"
    )

    score_formula = [
        "Pullback Score (30%): Cuánto ha retrocedido desde máximo reciente (oportunidad de entrada)",
        "Momentum Score (25%): Fuerza relativa vs S&P 500 en últimos 20 días",
        "Calidad Técnica (25%): Limpieza de la tendencia, sin gaps, sin crashes",
        "Sentimiento NLP (20%): Score de noticias recientes (positivo = catalizador)"
    ]
    for item in score_formula:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        "Se seleccionan las 5 acciones con mayor score de oportunidad."
    )

    doc.add_heading('7B.5 Gestión de la Posición', level=2)

    position_mgmt = [
        "Entrada: Al cierre del día de señal",
        "Stop Loss: 2% por debajo del precio de entrada",
        "Take Profit: Cierre automático al día 5 (holding period fijo)",
        "Salida anticipada: Si el stop loss es tocado intradiario"
    ]
    for item in position_mgmt:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7B.6 Ejemplo de Output del Sistema', level=2)

    doc.add_paragraph("El sistema genera un reporte como el siguiente:")

    doc.add_paragraph()

    # Simulación de output
    output_example = doc.add_paragraph()
    output_example.add_run("═══════════════════════════════════════════════════════════════\n").bold = True
    output_example.add_run("SWING TRADING ALERT - S&P 500 Top 5 (5 días)\n").bold = True
    output_example.add_run("Fecha: [Fecha actual] | Régimen: Bull Market | VIX: 18.5\n")
    output_example.add_run("═══════════════════════════════════════════════════════════════\n").bold = True
    output_example.add_run("\n")
    output_example.add_run("Rank | Ticker | Sector      | Score | RSI  | Pullback | Entry\n")
    output_example.add_run("─────────────────────────────────────────────────────────────\n")
    output_example.add_run("  1  | AAPL   | Technology  | 87.3  | 42   | -3.2%    | $185.20\n")
    output_example.add_run("  2  | MSFT   | Technology  | 84.1  | 45   | -2.8%    | $412.50\n")
    output_example.add_run("  3  | UNH    | Healthcare  | 81.7  | 38   | -4.1%    | $524.30\n")
    output_example.add_run("  4  | JPM    | Financials  | 79.4  | 44   | -2.5%    | $198.75\n")
    output_example.add_run("  5  | V      | Financials  | 77.8  | 41   | -3.0%    | $278.40\n")
    output_example.add_run("\n")
    output_example.add_run("Stop Loss: 2% | Horizonte: 5 días | Próximos earnings: Ninguno\n")
    output_example.add_run("═══════════════════════════════════════════════════════════════\n").bold = True

    doc.add_paragraph()

    doc.add_paragraph(
        "Este tipo de análisis se ejecuta diariamente de forma automática, permitiendo "
        "identificar oportunidades tácticas de corto plazo con alta probabilidad de éxito "
        "basadas en datos objetivos, no en intuición."
    )

    doc.add_page_break()

    # ========== 8. INFRAESTRUCTURA DE DATOS ==========
    doc.add_heading('8. Infraestructura de Datos', level=1)

    doc.add_paragraph(
        "El sistema está respaldado por una infraestructura de datos robusta y escalable, "
        "capaz de manejar desde el volumen actual hasta petabytes de información."
    )

    doc.add_heading('7.1 Volumen de Datos Actual', level=2)

    data_table = doc.add_table(rows=6, cols=3)
    data_table.style = 'Table Grid'

    data_info = [
        ('Categoría', 'Volumen', 'Cobertura'),
        ('Precios Históricos', '97 millones de registros', '86,000+ símbolos, 20+ años'),
        ('Fundamentales', '2.5 millones de registros', 'Balance sheets, income, cash flow'),
        ('Key Metrics & Ratios', '400,000 registros', 'P/E, ROE, márgenes, etc.'),
        ('Earnings Transcripts', '42,000 documentos', 'Transcripciones de llamadas'),
        ('ETF Holdings', '565,000 registros', '1,700+ ETFs con composición')
    ]

    for i, row_data in enumerate(data_info):
        for j, cell_text in enumerate(row_data):
            data_table.rows[i].cells[j].text = cell_text
            if i == 0:
                data_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    doc.add_heading('7.2 Actualización Continua', level=2)

    update_items = [
        "Precios: Actualización diaria automática tras cierre de mercado",
        "Fundamentales: Actualización semanal y tras earnings",
        "Noticias: Procesamiento en tiempo real con NLP",
        "Métricas Técnicas: Recálculo diario de 130+ indicadores"
    ]

    for item in update_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7.3 Calidad de Datos', level=2)

    doc.add_paragraph(
        "Se implementan múltiples capas de validación para garantizar la integridad de los datos:"
    )

    quality_items = [
        "Detección de outliers y valores anómalos",
        "Ajuste automático por splits y dividendos",
        "Validación cruzada entre fuentes",
        "Logs de auditoría para trazabilidad"
    ]

    for item in quality_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ========== 8. ROADMAP ==========
    doc.add_heading('9. Roadmap y Escalabilidad', level=1)

    doc.add_heading('9.1 Fases del Proyecto', level=2)

    roadmap_table = doc.add_table(rows=7, cols=4)
    roadmap_table.style = 'Table Grid'

    roadmap_data = [
        ('Fase', 'Descripción', 'Datos', 'Estado'),
        ('1', 'Infraestructura de datos y features técnicos', '35 GB', 'En progreso'),
        ('2', 'Features fundamentales y momentum', '50 GB', 'Estructura lista'),
        ('3', 'Módulo NLP y análisis de sentimiento', '100 GB', 'Arquitectura creada'),
        ('4', 'Modelos ML predictivos', '200 GB', 'Diseño completado'),
        ('5', 'Deep Learning y series temporales', '500 GB', 'Planificado'),
        ('6', 'Enterprise: Cluster distribuido', '1 TB+', 'Roadmap')
    ]

    for i, row_data in enumerate(roadmap_data):
        for j, cell_text in enumerate(row_data):
            roadmap_table.rows[i].cells[j].text = cell_text
            if i == 0:
                roadmap_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    doc.add_heading('9.2 Escalabilidad', level=2)

    doc.add_paragraph(
        "La arquitectura está diseñada para escalar horizontalmente. El sistema puede "
        "crecer de forma incremental sin reescribir código:"
    )

    scale_items = [
        "Base de datos: De PostgreSQL single-node a cluster distribuido",
        "Procesamiento: Soporte para procesamiento paralelo y batch",
        "Almacenamiento: Compatible con soluciones cloud (AWS, GCP, Azure)",
        "ML Models: Infraestructura para reentrenamiento automático"
    ]

    for item in scale_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ========== 10. ARQUITECTURA MACROECONÓMICA ==========
    doc.add_heading('10. Motor Macroeconómico', level=1)

    doc.add_paragraph(
        "El sistema incorpora un motor macroeconómico completo que analiza el contexto "
        "global del mercado para ajustar dinámicamente la estrategia de inversión."
    )

    doc.add_heading('10.1 Las 12 Categorías Macro', level=2)

    doc.add_paragraph(
        "Se monitorizan 12 dimensiones macroeconómicas, cada una con múltiples indicadores:"
    )

    # Tabla de categorías
    cat_table = doc.add_table(rows=13, cols=3)
    cat_table.style = 'Table Grid'

    cat_data = [
        ('#', 'Categoría', 'Indicadores Clave'),
        ('1', 'Inflación', 'CPI, PCE, PPI, salarios, expectativas'),
        ('2', 'Crecimiento', 'PIB, producción industrial, PMIs, consumo'),
        ('3', 'Laboral', 'Desempleo, payrolls, claims, vacantes'),
        ('4', 'Política Monetaria', 'Fed Funds, tipos reales, balance Fed'),
        ('5', 'Curva de Tipos', 'Yields 2Y/10Y/30Y, pendiente, inversión'),
        ('6', 'Liquidez', 'M2, crédito bancario, condiciones financieras'),
        ('7', 'Crédito', 'Spreads IG/HY, TED spread, defaults'),
        ('8', 'Estrés', 'VIX, MOVE, índices de estrés financiero'),
        ('9', 'Divisas', 'DXY, EUR/USD, balanza comercial'),
        ('10', 'Commodities', 'Petróleo, gas, cobre, índices'),
        ('11', 'Metales Preciosos', 'Oro, plata, ratio oro/tipos reales'),
        ('12', 'Sentimiento', 'EPU, confianza consumidor, geopolítico'),
    ]

    for i, row_data in enumerate(cat_data):
        for j, cell_text in enumerate(row_data):
            cat_table.rows[i].cells[j].text = cell_text
            if i == 0:
                cat_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    doc.add_heading('10.2 Los 7 Regímenes de Mercado', level=2)

    doc.add_paragraph(
        "En lugar de clasificaciones rígidas, el sistema identifica 7 regímenes naturales "
        "de mercado basados en la combinación de indicadores macro:"
    )

    # Tabla de regímenes
    regime_table = doc.add_table(rows=8, cols=3)
    regime_table.style = 'Table Grid'

    regime_data = [
        ('Régimen', 'Características', 'Estrategia'),
        ('Goldilocks', 'Crecimiento sólido, inflación baja, estrés bajo', 'Growth, Momentum, Tech'),
        ('Reflación', 'Crecimiento alto, inflación subiendo, commodities fuertes', 'Value, Cyclicals, Energy'),
        ('Tightening', 'Inflación alta, Fed hawkish, liquidez endureciendo', 'Defensivos, Quality, Healthcare'),
        ('Slowdown', 'Crecimiento cayendo, inflación enfriándose', 'Quality, Low Vol, selectivo'),
        ('Stagflation', 'Crecimiento débil + inflación alta', 'Commodities, Real Assets, Cash'),
        ('Credit Crisis', 'Estrés extremo, spreads disparados, deleveraging', 'Cash, Treasuries, Gold'),
        ('Recovery', 'Rebote post-crisis, mejora desde mínimos', 'Early Cyclicals, Small Caps, Financials'),
    ]

    for i, row_data in enumerate(regime_data):
        for j, cell_text in enumerate(row_data):
            regime_table.rows[i].cells[j].text = cell_text
            if i == 0:
                regime_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    doc.add_heading('10.3 Scores Compuestos', level=2)

    doc.add_paragraph(
        "Los múltiples indicadores se sintetizan en 5 scores dimensionales (-1 a +1):"
    )

    scores_items = [
        "Growth Score: Combina PIB, producción, PMIs, empleo → señala expansión/contracción",
        "Inflation Score: Combina CPI, PCE, expectativas → señala presión inflacionaria",
        "Stress Score: Combina VIX, spreads, NFCI → señala nivel de estrés sistémico",
        "Liquidity Score: Combina M2, crédito, lending standards → señala condiciones financieras",
        "Commodity Score: Combina energía, metales, índices → señala ciclo de materias primas"
    ]
    for item in scores_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('10.4 Clasificación de Indicadores', level=2)

    doc.add_paragraph(
        "Cada indicador está clasificado por su rol en el modelo:"
    )

    class_items = [
        "Leading: Anticipan el ciclo (PMIs, curva de tipos, claims) → señales tempranas",
        "Coincident: Describen el presente (producción, empleo) → confirmación",
        "Lagging: Confirman tarde (inflación core, desempleo) → validación",
        "Core Input: Definen régimen directamente (VIX, spreads, curva)",
        "Context Input: Explican pero no deciden (oro, dólar, sentiment)"
    ]
    for item in class_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('10.5 Ratios Diana por Régimen', level=2)

    doc.add_paragraph(
        "Los objetivos de cartera se ajustan automáticamente según el régimen detectado:"
    )

    # Tabla de ratios diana
    target_table = doc.add_table(rows=5, cols=5)
    target_table.style = 'Table Grid'

    target_data = [
        ('Parámetro', 'Goldilocks', 'Tightening', 'Crisis', 'Recovery'),
        ('P/E Máximo', '28', '18', '12', '25'),
        ('Beta Objetivo', '1.0-1.3', '0.6-0.9', '0.3-0.5', '1.2-1.5'),
        ('Cash %', '5%', '15%', '40%', '5%'),
        ('Defensivos %', '15%', '40%', '45%', '10%'),
    ]

    for i, row_data in enumerate(target_data):
        for j, cell_text in enumerate(row_data):
            target_table.rows[i].cells[j].text = cell_text
            if i == 0 or j == 0:
                target_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    doc.add_paragraph(
        "Esta integración macro-cuantitativa permite que la cartera se adapte automáticamente "
        "a las condiciones cambiantes del mercado, reduciendo drawdowns en periodos adversos "
        "y capturando oportunidades en expansiones."
    )

    doc.add_page_break()

    # ========== 11. VENTAJAS COMPETITIVAS ==========
    doc.add_heading('11. Ventajas Competitivas', level=1)

    advantages = [
        ("Enfoque Multi-Factor Integrado",
         "A diferencia de herramientas que analizan un solo aspecto, PatrimonioSmart "
         "integra técnico, fundamental, sentimiento y macro en un único sistema coherente."),

        ("Adaptabilidad a Regímenes",
         "El sistema no aplica reglas estáticas. Los parámetros se ajustan automáticamente "
         "según las condiciones del mercado, reduciendo drawdowns en periodos adversos."),

        ("Ratios Diana Dinámicos",
         "Los objetivos de valoración, riesgo y exposición cambian según el contexto, "
         "permitiendo ser más agresivo en bull markets y más defensivo en bear markets."),

        ("NLP para Información No Estructurada",
         "Capacidad de procesar y extraer señales de miles de noticias y transcripts "
         "de earnings, información que los sistemas tradicionales ignoran."),

        ("Escalabilidad Probada",
         "Arquitectura diseñada para crecer de 35 GB a 1 TB+ sin cambios estructurales, "
         "permitiendo incorporar nuevas fuentes de datos progresivamente."),

        ("Eliminación de Sesgos",
         "Las decisiones se basan en datos y reglas cuantitativas, eliminando los sesgos "
         "emocionales que afectan a la mayoría de inversores."),

        ("Transparencia y Auditabilidad",
         "Cada decisión está respaldada por datos y puede ser auditada. No hay 'cajas negras' "
         "incomprensibles.")
    ]

    for title, desc in advantages:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # ========== CIERRE ==========
    doc.add_paragraph()
    doc.add_paragraph()

    closing = doc.add_paragraph()
    closing_run = closing.add_run("PatrimonioSmart")
    closing_run.bold = True
    closing_run.font.size = Pt(16)
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tagline2 = doc.add_paragraph()
    tag2_run = tagline2.add_run("Inversión inteligente, adaptativa y basada en datos")
    tag2_run.italic = True
    tag2_run.font.size = Pt(12)
    tagline2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    contact = doc.add_paragraph()
    contact.add_run("Documento preparado para presentación a inversores")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_final = doc.add_paragraph()
    date_final.add_run(datetime.now().strftime("%d de %B de %Y"))
    date_final.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Guardar documento
    output_path = os.path.join(os.path.dirname(__file__), '..', 'PatrimonioSmart_Investor_Presentation.docx')
    doc.save(output_path)
    print(f"Documento generado: {os.path.abspath(output_path)}")
    return output_path

if __name__ == "__main__":
    create_investor_document()
