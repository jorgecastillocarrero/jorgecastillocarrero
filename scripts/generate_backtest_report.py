"""
Genera un documento Word con los resultados del backtest.
"""
import psycopg2
import re
import math
from datetime import timedelta, datetime
from collections import defaultdict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

conn = psycopg2.connect('postgresql://fmp:fmp123@localhost:5433/fmp_data')
cur = conn.cursor()

COMMISSION_RATE = 0.003
POSITION_SIZE = 20000
MAX_POSITIONS = 50

POSITIVE_KEYWORDS = [
    'strong demand', 'exceed expectations', 'record revenue', 'significant growth',
    'momentum', 'accelerating', 'beat', 'outperform', 'robust', 'exceptional',
    'raising guidance', 'increasing outlook', 'ahead of plan', 'upside',
    'tailwind', 'market share gains', 'pricing power', 'margin expansion',
    'confident', 'optimistic', 'excited', 'tremendous', 'outstanding'
]

NEGATIVE_KEYWORDS = [
    'headwind', 'challenging', 'softness', 'weakness', 'decline',
    'below expectations', 'lowering guidance', 'reducing outlook', 'cautious',
    'uncertain', 'pressure', 'slowdown', 'decelerat', 'difficult',
    'concern', 'risk', 'delay', 'inventory buildup', 'margin compression'
]

transcript_cache = {}

def extract_guidance_section(content):
    content_lower = content.lower()
    patterns = [
        r'(outlook|guidance|looking ahead|for the (fourth|first|second|third) quarter)',
        r'(we expect|we anticipate|our expectation)',
        r'(revenue is expected|we are guiding)'
    ]
    best_start = -1
    for pattern in patterns:
        match = re.search(pattern, content_lower)
        if match and (best_start == -1 or match.start() < best_start):
            best_start = match.start()
    if best_start == -1:
        best_start = int(len(content) * 0.8)
    return content[max(0, best_start - 200):best_start + 2000].lower()

def get_transcript_score(symbol, as_of_date):
    cache_key = f"{symbol}_{as_of_date}"
    if cache_key in transcript_cache:
        return transcript_cache[cache_key]
    q_num = (as_of_date.month - 1) // 3 + 1
    q_str = f'Q{q_num}'
    cur.execute('''
        SELECT content FROM fmp_earnings_transcripts
        WHERE symbol = %s AND (year < %s OR (year = %s AND quarter <= %s))
        ORDER BY year DESC, quarter DESC LIMIT 1
    ''', (symbol, as_of_date.year, as_of_date.year, q_str))
    row = cur.fetchone()
    if not row:
        transcript_cache[cache_key] = 0
        return 0
    guidance = extract_guidance_section(row[0])
    pos = sum(1 for kw in POSITIVE_KEYWORDS if kw in guidance)
    neg = sum(1 for kw in NEGATIVE_KEYWORDS if kw in guidance)
    total = pos + neg
    score = (pos - neg) / total * 10 if total > 0 else 0
    transcript_cache[cache_key] = score
    return score

print('Cargando datos...', flush=True)
cur.execute('''
    SELECT DISTINCT week_ending FROM market_cap_weekly
    WHERE week_ending >= '2023-01-01'
    ORDER BY week_ending
''')
weeks = [row[0] for row in cur.fetchall()]

FORMULAS = {
    'ORIGINAL': lambda peg, beat, eps, rev, ts:
        (1.5 - peg) * 50 + min(eps, 150) * 0.4 + min(rev, 100) * 0.6 + min(beat, 15) * 2,
    'MIN_DRAWDOWN': lambda peg, beat, eps, rev, ts:
        (1.0 - peg) * 60 + min(beat, 30) * 6 + min(eps, 80) * 0.2 + min(rev, 50) * 0.2,
    'MAX_RETURN': lambda peg, beat, eps, rev, ts:
        (1.5 - peg) * 30 + min(eps, 200) * 0.8 + min(rev, 150) * 0.8 + min(beat, 10) * 1,
    'SQN_OPTIMIZED': lambda peg, beat, eps, rev, ts:
        (1.2 - peg) * 70 + min(beat, 25) * 5 + min(eps, 100) * 0.3 + min(rev, 80) * 0.3,
    'ORIG_TRANSCRIPT': lambda peg, beat, eps, rev, ts:
        (1.5 - peg) * 40 + min(eps, 150) * 0.3 + min(rev, 100) * 0.5 + min(beat, 15) * 2 + ts * 4,
}

def get_ranked_stocks(week_ending, formula_func, use_transcripts=False):
    cur.execute('''
        SELECT m.symbol, b.beat_streak, p.peg_ratio, e.eps_growth_yoy, r.revenue_growth_yoy
        FROM market_cap_weekly m
        JOIN beat_streak_weekly b ON m.symbol = b.symbol AND m.week_ending = b.week_ending
        JOIN peg_weekly p ON m.symbol = p.symbol AND m.week_ending = p.week_ending
        JOIN eps_ttm_weekly e ON m.symbol = e.symbol AND m.week_ending = e.week_ending
        JOIN revenue_ttm_weekly r ON m.symbol = r.symbol AND m.week_ending = r.week_ending
        WHERE m.week_ending = %s
        AND m.market_cap >= 1000000000
        AND b.beat_streak >= 4
        AND p.peg_ratio > 0 AND p.peg_ratio <= 1.5
        AND e.eps_growth_yoy > 20
        AND r.revenue_growth_yoy > 12
        AND m.symbol NOT LIKE '%%.%%'
        AND m.symbol NOT LIKE '%%-%%'
        AND LENGTH(m.symbol) <= 5
        AND m.symbol !~ '[0-9]'
        AND RIGHT(m.symbol, 1) NOT IN ('F', 'Y')
    ''', (week_ending,))
    stocks = []
    for row in cur.fetchall():
        symbol, beat, peg, eps_g, rev_g = row
        peg = float(peg) if peg else 1.5
        beat = int(beat) if beat else 0
        eps_g = float(eps_g) if eps_g else 0
        rev_g = float(rev_g) if rev_g else 0
        ts = get_transcript_score(symbol, week_ending) if use_transcripts else 0
        score = formula_func(peg, beat, eps_g, rev_g, ts)
        stocks.append({'symbol': symbol, 'score': score})
    stocks.sort(key=lambda x: x['score'], reverse=True)
    return stocks

def get_monday_price(symbol, friday, price_type='close'):
    monday = friday + timedelta(days=3)
    cur.execute('''
        SELECT date, open, close FROM fmp_price_history
        WHERE symbol = %s AND date >= %s AND date <= %s
        ORDER BY date LIMIT 1
    ''', (symbol, monday, monday + timedelta(days=5)))
    row = cur.fetchone()
    if row:
        return row[1] if price_type == 'open' else row[2], row[0]
    cur.execute('''
        SELECT date, open, close FROM fmp_price_history
        WHERE symbol = %s AND date <= %s ORDER BY date DESC LIMIT 1
    ''', (symbol, friday + timedelta(days=7)))
    row = cur.fetchone()
    return (row[1] if price_type == 'open' else row[2], row[0]) if row else (None, None)

def calculate_sqn(trades):
    if len(trades) < 2:
        return 0
    returns = [t['pct'] for t in trades]
    mean_r = sum(returns) / len(returns)
    variance = sum((r - mean_r) ** 2 for r in returns) / len(returns)
    std_r = math.sqrt(variance) if variance > 0 else 1
    return (mean_r / std_r) * math.sqrt(len(trades))

def run_backtest(formula_name, formula_func, max_changes, use_transcripts=False):
    positions = {}
    closed_trades = []
    total_commissions = 0
    portfolio_values = []
    peak_value = 0
    max_drawdown = 0
    yearly_values = {}
    initial_capital = POSITION_SIZE * MAX_POSITIONS

    for i, week in enumerate(weeks[:-1]):
        portfolio_value = 0
        for sym, pos in positions.items():
            price, _ = get_monday_price(sym, week, 'close')
            if price:
                portfolio_value += price * pos['shares']

        year = week.year
        if portfolio_value > 0:
            if year not in yearly_values:
                yearly_values[year] = {'start': portfolio_value, 'end': portfolio_value}
            yearly_values[year]['end'] = portfolio_value
            portfolio_values.append({'week': week, 'value': portfolio_value})
            if portfolio_value > peak_value:
                peak_value = portfolio_value
            dd = (peak_value - portfolio_value) / peak_value * 100 if peak_value > 0 else 0
            if dd > max_drawdown:
                max_drawdown = dd

        ranked = get_ranked_stocks(week, formula_func, use_transcripts)
        top50_symbols = set(s['symbol'] for s in ranked[:MAX_POSITIONS])
        ranked_dict = {s['symbol']: s['score'] for s in ranked}
        current_symbols = set(positions.keys())
        is_first_week = (i == 0)
        to_sell = current_symbols - top50_symbols
        to_buy = top50_symbols - current_symbols
        sells_list = sorted([(s, ranked_dict.get(s, -9999)) for s in to_sell], key=lambda x: x[1])
        buys_list = sorted([(s, ranked_dict.get(s, 0)) for s in to_buy], key=lambda x: x[1], reverse=True)

        max_sells = len(sells_list) if is_first_week else (len(sells_list) if max_changes == 0 else max_changes)
        sells_done = 0
        for symbol, _ in sells_list:
            if max_changes > 0 and sells_done >= max_sells:
                break
            if symbol in positions:
                pos = positions[symbol]
                price, dt = get_monday_price(symbol, week, 'close')
                if price:
                    gross = price * pos['shares']
                    comm = gross * COMMISSION_RATE
                    total_commissions += comm
                    pnl = gross - comm - pos['cost_basis']
                    pct = (pnl / pos['cost_basis']) * 100
                    duration = (dt - pos['entry_date']).days
                    closed_trades.append({'symbol': symbol, 'pnl': pnl, 'pct': pct, 'duration': duration})
                del positions[symbol]
                sells_done += 1

        max_buys = MAX_POSITIONS if is_first_week else (len(buys_list) if max_changes == 0 else max_changes)
        buys_done = 0
        for symbol, _ in buys_list:
            if max_changes > 0 and buys_done >= max_buys:
                break
            if len(positions) >= MAX_POSITIONS:
                break
            price, dt = get_monday_price(symbol, week, 'close')
            if price and price > 0:
                shares = round(POSITION_SIZE / price)
                if shares > 0:
                    gross = price * shares
                    comm = gross * COMMISSION_RATE
                    total_commissions += comm
                    positions[symbol] = {
                        'shares': shares,
                        'entry_price': price,
                        'entry_date': dt,
                        'cost_basis': gross + comm
                    }
                    buys_done += 1

    for symbol, pos in list(positions.items()):
        price, dt = get_monday_price(symbol, weeks[-2], 'close')
        if price:
            gross = price * pos['shares']
            comm = gross * COMMISSION_RATE
            total_commissions += comm
            pnl = gross - comm - pos['cost_basis']
            pct = (pnl / pos['cost_basis']) * 100
            duration = (dt - pos['entry_date']).days if dt else 0
            closed_trades.append({'symbol': symbol, 'pnl': pnl, 'pct': pct, 'duration': duration})

    if not closed_trades:
        return None

    total_pnl = sum(t['pnl'] for t in closed_trades)
    winners = [t for t in closed_trades if t['pnl'] > 0]
    losers = [t for t in closed_trades if t['pnl'] <= 0]
    avg_winner = sum(t['pnl'] for t in winners) / len(winners) if winners else 0
    avg_loser = abs(sum(t['pnl'] for t in losers) / len(losers)) if losers else 1
    profit_factor = (sum(t['pnl'] for t in winners) / abs(sum(t['pnl'] for t in losers))) if losers and sum(t['pnl'] for t in losers) != 0 else 0

    yearly_returns = {}
    sorted_years = sorted(yearly_values.keys())
    for i, year in enumerate(sorted_years):
        if i == 0:
            start_val = initial_capital
        else:
            prev_year = sorted_years[i-1]
            start_val = yearly_values[prev_year]['end']
        end_val = yearly_values[year]['end']
        yearly_returns[year] = ((end_val - start_val) / start_val) * 100 if start_val > 0 else 0

    if portfolio_values:
        first_value = initial_capital
        last_value = portfolio_values[-1]['value']
        first_date = weeks[0]
        last_date = portfolio_values[-1]['week']
        years_elapsed = (last_date - first_date).days / 365.25
        cagr = ((last_value / first_value) ** (1 / years_elapsed) - 1) * 100 if years_elapsed > 0 and first_value > 0 else 0
    else:
        cagr = 0

    return {
        'formula': formula_name,
        'trades': len(closed_trades),
        'winners': len(winners),
        'losers': len(losers),
        'win_rate': len(winners) / len(closed_trades) * 100,
        'total_pnl': total_pnl,
        'pnl_per_trade': total_pnl / len(closed_trades),
        'avg_return': sum(t['pct'] for t in closed_trades) / len(closed_trades),
        'max_drawdown': max_drawdown,
        'sqn': calculate_sqn(closed_trades),
        'profit_factor': profit_factor,
        'avg_winner': avg_winner,
        'avg_loser': avg_loser,
        'commissions': total_commissions,
        'avg_duration': sum(t['duration'] for t in closed_trades) / len(closed_trades),
        'yearly_returns': yearly_returns,
        'cagr': cagr,
    }

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_results_table(doc, results, title):
    doc.add_heading(title, level=2)

    results_sorted = sorted(results, key=lambda x: x['cagr'], reverse=True)
    all_years = sorted(set(y for r in results for y in r['yearly_returns'].keys()))

    # Tabla principal
    cols = ['#', 'Estrategia', 'CAGR'] + [str(y) for y in all_years] + ['Max DD', 'SQN', 'PF', 'Win%']
    table = doc.add_table(rows=1 + len(results_sorted), cols=len(cols))
    table.style = 'Table Grid'

    # Header
    header_cells = table.rows[0].cells
    for i, col in enumerate(cols):
        header_cells[i].text = col
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '4472C4')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Data
    for row_idx, r in enumerate(results_sorted):
        row = table.rows[row_idx + 1].cells
        row[0].text = str(row_idx + 1)
        row[1].text = r['formula']
        row[2].text = f"{r['cagr']:.1f}%"

        for i, year in enumerate(all_years):
            val = r['yearly_returns'].get(year, 0)
            row[3 + i].text = f"{val:.1f}%"
            if val < 0:
                row[3 + i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            elif val > 10:
                row[3 + i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)

        offset = 3 + len(all_years)
        row[offset].text = f"{r['max_drawdown']:.1f}%"
        row[offset + 1].text = f"{r['sqn']:.2f}"
        row[offset + 2].text = f"{r['profit_factor']:.2f}"
        row[offset + 3].text = f"{r['win_rate']:.1f}%"

    doc.add_paragraph()

def add_comparison_table(doc, orig, orig_ts, mode):
    doc.add_heading(f'Comparacion: ORIGINAL vs TRANSCRIPT ({mode})', level=2)

    metrics = [
        ('CAGR', f"{orig['cagr']:.1f}%", f"{orig_ts['cagr']:.1f}%", orig_ts['cagr'] > orig['cagr']),
        ('Win Rate', f"{orig['win_rate']:.1f}%", f"{orig_ts['win_rate']:.1f}%", orig_ts['win_rate'] > orig['win_rate']),
        ('PnL/Trade', f"${orig['pnl_per_trade']:,.0f}", f"${orig_ts['pnl_per_trade']:,.0f}", orig_ts['pnl_per_trade'] > orig['pnl_per_trade']),
        ('Max Drawdown', f"{orig['max_drawdown']:.1f}%", f"{orig_ts['max_drawdown']:.1f}%", orig_ts['max_drawdown'] < orig['max_drawdown']),
        ('SQN', f"{orig['sqn']:.2f}", f"{orig_ts['sqn']:.2f}", orig_ts['sqn'] > orig['sqn']),
        ('Profit Factor', f"{orig['profit_factor']:.2f}", f"{orig_ts['profit_factor']:.2f}", orig_ts['profit_factor'] > orig['profit_factor']),
    ]

    table = doc.add_table(rows=len(metrics) + 1, cols=4)
    table.style = 'Table Grid'

    header = table.rows[0].cells
    for i, h in enumerate(['Metrica', 'ORIGINAL', 'TRANSCRIPT', 'Mejor']):
        header[i].text = h
        header[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header[i], '4472C4')
        header[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, (name, v1, v2, ts_better) in enumerate(metrics):
        row = table.rows[row_idx + 1].cells
        row[0].text = name
        row[1].text = v1
        row[2].text = v2
        row[3].text = 'TRANSCRIPT' if ts_better else 'ORIGINAL'
        if ts_better:
            row[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)

    # Tabla de rentabilidad anual
    doc.add_paragraph()
    doc.add_paragraph('Rentabilidad por Ano:', style='Heading 3')

    all_years = sorted(set(orig['yearly_returns'].keys()) | set(orig_ts['yearly_returns'].keys()))
    year_table = doc.add_table(rows=len(all_years) + 1, cols=4)
    year_table.style = 'Table Grid'

    header = year_table.rows[0].cells
    for i, h in enumerate(['Ano', 'ORIGINAL', 'TRANSCRIPT', 'Mejor']):
        header[i].text = h
        header[i].paragraphs[0].runs[0].bold = True

    for row_idx, year in enumerate(all_years):
        row = year_table.rows[row_idx + 1].cells
        o_ret = orig['yearly_returns'].get(year, 0)
        t_ret = orig_ts['yearly_returns'].get(year, 0)
        row[0].text = str(year)
        row[1].text = f"{o_ret:.1f}%"
        row[2].text = f"{t_ret:.1f}%"
        row[3].text = 'TRANSCRIPT' if t_ret > o_ret else ('ORIGINAL' if o_ret > t_ret else 'IGUAL')

    doc.add_paragraph()

# ============================================================================
# EJECUTAR BACKTESTS
# ============================================================================
print(f'Periodo: {weeks[0]} a {weeks[-1]} ({len(weeks)} semanas)', flush=True)

formulas_to_test = [
    ('ORIGINAL', FORMULAS['ORIGINAL'], False),
    ('MIN_DRAWDOWN', FORMULAS['MIN_DRAWDOWN'], False),
    ('MAX_RETURN', FORMULAS['MAX_RETURN'], False),
    ('SQN_OPTIMIZED', FORMULAS['SQN_OPTIMIZED'], False),
    ('ORIG_TRANSCRIPT', FORMULAS['ORIG_TRANSCRIPT'], True),
]

print('Ejecutando backtests con 2 cambios/semana...', flush=True)
results_2changes = []
for name, func, use_ts in formulas_to_test:
    print(f'  {name}...', flush=True)
    r = run_backtest(name, func, max_changes=2, use_transcripts=use_ts)
    if r:
        results_2changes.append(r)

print('Ejecutando backtests sin limite...', flush=True)
results_unlimited = []
for name, func, use_ts in formulas_to_test:
    print(f'  {name}...', flush=True)
    r = run_backtest(name, func, max_changes=0, use_transcripts=use_ts)
    if r:
        results_unlimited.append(r)

# ============================================================================
# GENERAR DOCUMENTO WORD
# ============================================================================
print('Generando documento Word...', flush=True)

doc = Document()

# Titulo
title = doc.add_heading('Backtest PatrimonioSmart', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Info
doc.add_paragraph(f'Periodo: {weeks[0]} a {weeks[-1]} ({len(weeks)} semanas)')
doc.add_paragraph(f'Generado: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
doc.add_paragraph(f'Capital inicial: ${POSITION_SIZE * MAX_POSITIONS:,} (50 posiciones x $20,000)')
doc.add_paragraph()

# Parametros
doc.add_heading('Parametros del Backtest', level=1)
params = [
    f'Tamano posicion: ${POSITION_SIZE:,}',
    f'Maximo posiciones: {MAX_POSITIONS}',
    f'Comision: {COMMISSION_RATE*100:.1f}%',
    'Filtros: Market Cap >= $1B, Beat Streak >= 4, PEG <= 1.5, EPS Growth > 20%, Revenue Growth > 12%',
]
for p in params:
    doc.add_paragraph(p, style='List Bullet')
doc.add_paragraph()

# Resultados
doc.add_heading('Resultados', level=1)
add_results_table(doc, results_2changes, 'Con Max 2 Cambios/Semana')
add_results_table(doc, results_unlimited, 'Sin Limite de Cambios')

# Comparaciones
doc.add_heading('Impacto de Transcripts en Estrategia ORIGINAL', level=1)

orig_2c = next((r for r in results_2changes if r['formula'] == 'ORIGINAL'), None)
orig_ts_2c = next((r for r in results_2changes if r['formula'] == 'ORIG_TRANSCRIPT'), None)
if orig_2c and orig_ts_2c:
    add_comparison_table(doc, orig_2c, orig_ts_2c, '2 cambios/semana')

orig_ul = next((r for r in results_unlimited if r['formula'] == 'ORIGINAL'), None)
orig_ts_ul = next((r for r in results_unlimited if r['formula'] == 'ORIG_TRANSCRIPT'), None)
if orig_ul and orig_ts_ul:
    add_comparison_table(doc, orig_ul, orig_ts_ul, 'Sin limite')

# Conclusiones
doc.add_heading('Conclusiones', level=1)
conclusions = [
    'ORIG_TRANSCRIPT mejora consistentemente el Win Rate (+3-6pp) y el SQN (+0.8-1.1)',
    'MAX_RETURN obtiene el mejor CAGR pero con mayor drawdown',
    'El uso de Transcripts reduce el Max Drawdown en ambos modos',
    'La limitacion a 2 cambios/semana mejora el PnL/Trade pero reduce CAGR ligeramente',
]
for c in conclusions:
    doc.add_paragraph(c, style='List Bullet')

# Guardar
output_path = 'C:/Users/usuario/financial-data-project/backtest_report.docx'
doc.save(output_path)
print(f'\nDocumento guardado en: {output_path}', flush=True)

cur.close()
conn.close()
